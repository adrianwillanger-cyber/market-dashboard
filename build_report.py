from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page setup: US Letter, 1-inch margins ──────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(1)
section.right_margin  = Inches(1)
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)

# ── Helper: shade a table cell ─────────────────────────────────────────────
def shade_cell(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge, color in kwargs.items():
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'),   'single')
        tag.set(qn('w:sz'),    '4')
        tag.set(qn('w:color'), color)
        tcBorders.append(tag)
    tcPr.append(tcBorders)

# ── Color palette ──────────────────────────────────────────────────────────
DARK_BLUE   = '1F3864'
MID_BLUE    = '2E5B8A'
ACCENT_BLUE = 'D5E8F4'
LIGHT_GRAY  = 'F2F2F2'
WHITE       = 'FFFFFF'
TEXT_DARK   = '1A1A1A'

# ══════════════════════════════════════════════════════════════════════════
# COVER BLOCK
# ══════════════════════════════════════════════════════════════════════════
# Solid dark banner via a 1-row, 1-col table
banner = doc.add_table(rows=1, cols=1)
banner.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = banner.cell(0, 0)
shade_cell(cell, DARK_BLUE)
cell.width = Inches(6.5)

p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
p.paragraph_format.space_after  = Pt(4)
run = p.add_run('OWNER PROPERTY RECORD')
run.bold      = True
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
run.font.name = 'Arial'

p2 = cell.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_before = Pt(0)
p2.paragraph_format.space_after  = Pt(6)
run2 = p2.add_run('57 Crooked Lane — Bonner County, Idaho')
run2.bold      = True
run2.font.size = Pt(16)
run2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
run2.font.name = 'Arial'

p3 = cell.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_before = Pt(0)
p3.paragraph_format.space_after  = Pt(6)
run3 = p3.add_run('Underground Fire Suppression System & CC&R Research Report')
run3.font.size = Pt(11)
run3.font.color.rgb = RGBColor(0xCC, 0xDD, 0xEE)
run3.font.name = 'Arial'
run3.italic = True

p4 = cell.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
p4.paragraph_format.space_before = Pt(4)
p4.paragraph_format.space_after  = Pt(16)
run4 = p4.add_run('Prepared: May 4, 2026  |  Parcel: RP037080000110A')
run4.font.size = Pt(9)
run4.font.color.rgb = RGBColor(0xAA, 0xBB, 0xCC)
run4.font.name = 'Arial'

doc.add_paragraph()  # spacer

# ══════════════════════════════════════════════════════════════════════════
# SUMMARY SNAPSHOT TABLE
# ══════════════════════════════════════════════════════════════════════════
snap = doc.add_table(rows=6, cols=2)
snap.alignment = WD_TABLE_ALIGNMENT.CENTER
col_w = [Inches(2.4), Inches(4.1)]

rows_data = [
    ('Property Address',   '57 Crooked Lane, Sandpoint, Idaho'),
    ('Parcel Number',      'RP037080000110A'),
    ('County',             'Bonner County, Idaho'),
    ('Lot Size',           '1.72 acres'),
    ('Year Built',         '2000'),
    ('Subdivision Profile','~10 lots, avg. 1.09 ac. — rural, outside municipal water service'),
]

for i, (label, value) in enumerate(rows_data):
    row = snap.rows[i]
    for j, cell in enumerate(row.cells):
        cell.width = col_w[j]
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        p.paragraph_format.left_indent  = Pt(6)
        if j == 0:
            shade_cell(cell, ACCENT_BLUE)
            run = p.add_run(label)
            run.bold = True
            run.font.size = Pt(9.5)
            run.font.name = 'Arial'
            run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        else:
            shade_cell(cell, WHITE if i % 2 == 0 else LIGHT_GRAY)
            run = p.add_run(value)
            run.font.size = Pt(9.5)
            run.font.name = 'Arial'
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

doc.add_paragraph()  # spacer

# ══════════════════════════════════════════════════════════════════════════
# HELPER: section heading
# ══════════════════════════════════════════════════════════════════════════
def add_section_heading(text, number):
    # Colored rule above
    p_rule = doc.add_paragraph()
    p_rule.paragraph_format.space_before = Pt(14)
    p_rule.paragraph_format.space_after  = Pt(0)
    pPr = p_rule._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top  = OxmlElement('w:bottom')
    top.set(qn('w:val'),   'single')
    top.set(qn('w:sz'),    '6')
    top.set(qn('w:color'), MID_BLUE)
    pBdr.append(top)
    pPr.append(pBdr)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)

    num_run = p.add_run(f'{number}  ')
    num_run.bold = True
    num_run.font.size = Pt(13)
    num_run.font.name = 'Arial'
    num_run.font.color.rgb = RGBColor(0x2E, 0x5B, 0x8A)

    txt_run = p.add_run(text)
    txt_run.bold = True
    txt_run.font.size = Pt(13)
    txt_run.font.name = 'Arial'
    txt_run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

def add_body(text, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Inches(0.3)
    if bold_prefix:
        br = p.add_run(bold_prefix)
        br.bold = True
        br.font.size = Pt(10)
        br.font.name = 'Arial'
        br.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

def add_callout(text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    shade_cell(cell, 'FFF3CD')
    set_cell_border(cell, left='E6A817')
    cp = cell.paragraphs[0]
    cp.paragraph_format.space_before = Pt(6)
    cp.paragraph_format.space_after  = Pt(6)
    cp.paragraph_format.left_indent  = Pt(6)
    run = cp.add_run(text)
    run.font.size = Pt(9.5)
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(0x5A, 0x3E, 0x00)
    run.italic = True
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 1 — Property Overview
# ══════════════════════════════════════════════════════════════════════════
add_section_heading('Property Overview', '1.')
add_body(
    'The subject property at 57 Crooked Lane is a 1.72-acre parcel in Bonner County, Idaho '
    '(parcel RP037080000110A), built in 2000. It sits within a small rural subdivision of '
    'approximately ten lots averaging 1.09 acres each. The subdivision lies outside city water '
    'service boundaries — a fact that directly drives the regulatory requirement for a shared '
    'underground fire suppression cistern serving the neighborhood.'
)

# ══════════════════════════════════════════════════════════════════════════
# SECTION 2 — Regulatory Basis for the Tank
# ══════════════════════════════════════════════════════════════════════════
add_section_heading('Regulatory Basis for the Underground Fire Suppression Tank', '2.')
add_body(
    'Bonner County Code (Title 12 — Land Use Regulations) explicitly requires that "adequate water '
    'supplies for fire suppression, as determined by the applicable fire district chief, shall be '
    'provided" in all subdivisions. For rural and semi-rural subdivisions without access to a '
    'municipal hydrant system — precisely the profile of Crooked Lane — developers are required '
    'to install a shared underground cistern. The applicable standards governing design and '
    'construction are:'
)
add_bullet(
    'NFPA 22 — Standard for Water Tanks for Private Fire Protection', '')
add_bullet(
    'NFPA 1142 — Water Supplies for Suburban and Rural Fire Fighting', '')
add_body(
    'Minimum capacity under these standards is typically 10,000 gallons. Actual installed '
    'capacity on Crooked Lane may range from 10,000 to 25,000 gallons depending on the original '
    'fire district determination at the time of subdivision platting.',
    space_after=8
)

# ══════════════════════════════════════════════════════════════════════════
# SECTION 3 — Tank Details & Physical Description
# ══════════════════════════════════════════════════════════════════════════
add_section_heading('Tank Physical Description', '3.')
add_body(
    'Based on standard practice for Bonner County rural subdivisions of this era, the tank is '
    'most likely a precast concrete or fiberglass underground cistern. Key physical characteristics '
    'typically include:'
)
add_bullet('Capacity: 10,000–25,000 gallons', '')
add_bullet('Construction: Precast concrete or fiberglass, fully underground', '')
add_bullet(
    'Access point: A dry hydrant standpipe at the road surface that fire trucks connect to directly, '
    'enabling drafting without pressurized water supply', '')
add_bullet('Location: Near the road edge, within or adjacent to the subdivision easement area', '')

add_body('')
add_callout(
    '⚠  Regulatory note: Idaho DEQ does not regulate fire-suppression-only water tanks. '
    'DEQ jurisdiction covers petroleum storage tanks only (Idaho Code § 41-4901 et seq.). '
    'There is no state registry for these cisterns — all records are held locally by the county '
    'recorder and the applicable fire district.'
)

# ══════════════════════════════════════════════════════════════════════════
# SECTION 4 — Maintenance & Ownership Responsibility
# ══════════════════════════════════════════════════════════════════════════
add_section_heading('Maintenance & Ownership Responsibility', '4.')
add_body(
    'This is the most operationally significant aspect of the tank for owners. Responsibility '
    'for the cistern does not fall on Bonner County, the city, or the fire district. Maintenance '
    'obligations rest with the homeowners of the subdivision collectively, almost certainly as '
    'spelled out in the recorded CC&Rs. Key points:'
)
add_bullet(
    'The CC&Rs govern cost-sharing for tank inspection, maintenance, and repair among all '
    'lot owners in the subdivision.', '')
add_bullet(
    'Idaho fire districts are required to inspect and test the tank periodically. Property '
    'owners and HOAs must provide access for these inspections.', '')
add_bullet(
    'If the HOA has gone dormant over the years (common on small rural subdivisions), it is '
    'critical to determine who is currently responsible for monitoring the tank.', '')
add_bullet(
    'Given elevated wildfire risk in the North Idaho corridor, a non-functioning or '
    'inadequately maintained cistern could constitute a liability issue — and may be flagged '
    'as a material defect at resale.', '')

doc.add_paragraph()
add_callout(
    '⚠  Action item: Confirm current tank condition, last inspection date, and active HOA '
    'status before any sale, refinancing, or insurance renewal. Request records from the '
    'applicable fire district and/or Bonner County.'
)

# ══════════════════════════════════════════════════════════════════════════
# SECTION 5 — Locating the CC&Rs
# ══════════════════════════════════════════════════════════════════════════
add_section_heading('How to Locate the Recorded CC&Rs', '5.')
add_body(
    'The Declaration of Covenants, Conditions & Restrictions for the Crooked Lane subdivision '
    'is recorded with the Bonner County Recorder and attached to the original subdivision plat. '
    'Four practical paths to obtain a copy are outlined below.'
)

steps = [
    ('Bonner County Property Search Portal',
     'Search parcel RP037080000110A at cloudgisapps.bonnercountyid.gov/PropertySearch/ to pull '
     'all recorded documents associated with the lot, including the original plat and any '
     'recorded CC&Rs.'),
    ('Bonner County Recorder (phone)',
     'Call (208) 265-1432. Reference your parcel number and ask for the recorded Declaration of '
     'Covenants, Conditions & Restrictions for the Crooked Lane subdivision. Copies are available '
     'for a small per-page fee and can be mailed or emailed.'),
    ('Original Closing/Title Documents',
     'CC&Rs are always listed as Schedule B exceptions in a title commitment. If you have the '
     'closing binder from original purchase, the CC&Rs are included there.'),
    ('Idaho Secretary of State',
     'If the HOA is registered as a nonprofit entity in Idaho, the registered agent and entity '
     'status can be found at sos.idaho.gov. This can help identify who has current HOA authority.'),
]

for i, (title, desc) in enumerate(steps, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.1)
    num_run = p.add_run(f'Step {i}: ')
    num_run.bold = True
    num_run.font.size = Pt(10)
    num_run.font.name = 'Arial'
    num_run.font.color.rgb = RGBColor(0x2E, 0x5B, 0x8A)
    title_run = p.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(10)
    title_run.font.name = 'Arial'
    title_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    dp = doc.add_paragraph()
    dp.paragraph_format.space_before = Pt(0)
    dp.paragraph_format.space_after  = Pt(6)
    dp.paragraph_format.left_indent  = Inches(0.3)
    dr = dp.add_run(desc)
    dr.font.size = Pt(10)
    dr.font.name = 'Arial'
    dr.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# ══════════════════════════════════════════════════════════════════════════
# SECTION 6 — Key Standards & References
# ══════════════════════════════════════════════════════════════════════════
add_section_heading('Applicable Standards & References', '6.')

refs = [
    ('Bonner County Code, Title 12',      'Land Use Regulations — fire suppression water supply requirements for subdivisions'),
    ('NFPA 22',                            'Standard for Water Tanks for Private Fire Protection'),
    ('NFPA 1142',                          'Water Supplies for Suburban and Rural Fire Fighting'),
    ('Idaho DEQ — UST Program',            'Confirms fire suppression tanks are not regulated under state UST rules (petroleum only)'),
    ('Bonner County Recorder',             '(208) 265-1432 — source for recorded CC&Rs and plat documents'),
    ('Bonner County Property Search',      'cloudgisapps.bonnercountyid.gov/PropertySearch/ — parcel RP037080000110A'),
    ('Idaho Secretary of State',           'sos.idaho.gov — HOA nonprofit registration lookup'),
]

ref_tbl = doc.add_table(rows=len(refs), cols=2)
ref_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
ref_col_w = [Inches(2.2), Inches(4.3)]

for i, (source, desc) in enumerate(refs):
    row = ref_tbl.rows[i]
    row.cells[0].width = ref_col_w[0]
    row.cells[1].width = ref_col_w[1]

    shade_cell(row.cells[0], ACCENT_BLUE if i % 2 == 0 else 'E8F2FB')
    shade_cell(row.cells[1], WHITE if i % 2 == 0 else LIGHT_GRAY)

    p0 = row.cells[0].paragraphs[0]
    p0.paragraph_format.space_before = Pt(4)
    p0.paragraph_format.space_after  = Pt(4)
    p0.paragraph_format.left_indent  = Pt(5)
    r0 = p0.add_run(source)
    r0.bold = True
    r0.font.size = Pt(9)
    r0.font.name = 'Arial'
    r0.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    p1 = row.cells[1].paragraphs[0]
    p1.paragraph_format.space_before = Pt(4)
    p1.paragraph_format.space_after  = Pt(4)
    p1.paragraph_format.left_indent  = Pt(5)
    r1 = p1.add_run(desc)
    r1.font.size = Pt(9)
    r1.font.name = 'Arial'
    r1.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

# ══════════════════════════════════════════════════════════════════════════
# FOOTER note
# ══════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
footer_rule = doc.add_paragraph()
footer_rule.paragraph_format.space_before = Pt(10)
footer_rule.paragraph_format.space_after  = Pt(4)
pPr2 = footer_rule._p.get_or_add_pPr()
pBdr2 = OxmlElement('w:pBdr')
top2  = OxmlElement('w:top')
top2.set(qn('w:val'),   'single')
top2.set(qn('w:sz'),    '4')
top2.set(qn('w:color'), 'AAAAAA')
pBdr2.append(top2)
pPr2.append(pBdr2)

fp = doc.add_paragraph()
fp.paragraph_format.space_before = Pt(0)
fp.paragraph_format.space_after  = Pt(0)
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run(
    'This document is prepared for owner records purposes. Contents are based on publicly '
    'available research as of May 2026. Consult a licensed attorney, title company, or '
    'professional engineer for legal or engineering determinations.'
)
fr.font.size = Pt(8)
fr.font.name = 'Arial'
fr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
fr.italic = True

# ══════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════
out_path = '/sessions/relaxed-focused-pascal/mnt/outputs/57_Crooked_Lane_Owner_Record.docx'
doc.save(out_path)
print(f'Saved: {out_path}')
